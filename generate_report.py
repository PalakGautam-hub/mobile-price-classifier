"""Generate the project report as a .docx file."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ---------------------------------------------------------------------------
# Global style defaults
# ---------------------------------------------------------------------------
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Times New Roman"
    hs.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(text, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5 + level * 0.25)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p


def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    doc.add_paragraph()  # spacer
    return table


# ===========================  1. TITLE PAGE  ===============================
doc.add_paragraph()
doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tp.add_run("Mobile Price Range Classifier\nUsing Machine Learning and Flask")
run.bold = True
run.font.size = Pt(24)
run.font.name = "Times New Roman"
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

add_para("Course Name: Application Development Laboratory (CS33002)",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para("Submitted By:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

members = [
    ("23053124", "Divyanshi Chaurasia"),
    ("23052982", "Palak Gautam"),
    ("23052547", "Anshu Shahdeo"),
    ("23052976", "Shabbir Uddin"),
    ("23052344", "Rahul Mishra"),
    ("23053529", "Animesh Krishnan"),
]
make_table(["Roll No", "Name"], members)

add_para("Submitted To: Mr. Anirban Saha (Course Instructor)",
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()
add_para("Department of Computer Science and Engineering",
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("KIIT University", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("6th Semester | Spring 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ==================  2. GROUP MEMBER CONTRIBUTION  =========================
add_heading("1. Group Member Contribution", level=1)
add_para("The contribution of each group member is detailed below:")

contributions = [
    ("23053124", "Divyanshi Chaurasia",
     "Dataset collection and preparation; Data preprocessing and feature engineering"),
    ("23052982", "Palak Gautam",
     "Model training and evaluation; Comparative analysis of ML models"),
    ("23052547", "Anshu Shahdeo",
     "HTML page design; CSS styling and responsive UI"),
    ("23052976", "Shabbir Uddin",
     "Flask backend development; Route handling and API integration"),
    ("23052344", "Rahul Mishra",
     "Testing and debugging; Frontend-backend communication"),
    ("23053529", "Animesh Krishnan",
     "Documentation and report preparation; Presentation and demonstration"),
]
make_table(["Roll No", "Name", "Responsibilities"], contributions)

doc.add_page_break()

# =====================  3. TABLE OF CONTENTS  ==============================
add_heading("2. Table of Contents", level=1)

toc_items = [
    "1. Group Member Contribution",
    "2. Table of Contents",
    "3. List of Figures",
    "4. List of Tables",
    "5. Introduction",
    "6. Problem Statement",
    "7. Objectives of the Project",
    "8. Dataset Description",
    "9. Data Preprocessing",
    "10. Methodology",
    "11. Models Tried for the Project",
    "12. Model Comparison and Best Model Selection",
    "13. Implementation Details",
    "14. Frontend Design",
    "15. Flask Integration",
    "16. Results and Discussion",
    "17. Testing",
    "18. Challenges Faced",
    "19. Conclusion",
]
for item in toc_items:
    add_para(item, space_after=2)

doc.add_page_break()

# =====================  4. LIST OF FIGURES  ================================
add_heading("3. List of Figures", level=1)
figures_list = [
    "Figure 1: System Architecture Diagram",
    "Figure 2: Main Web Interface with Input Form",
    "Figure 3: Input Validation Error Display",
    "Figure 4: Prediction Result — Low Price",
    "Figure 5: Prediction Result — Very High Price",
    "Figure 6: Confusion Matrix — Random Forest",
    "Figure 7: Model Accuracy Comparison Bar Chart",
]
for f in figures_list:
    add_para(f, space_after=2)

doc.add_paragraph()

# =====================  5. LIST OF TABLES  =================================
add_heading("4. List of Tables", level=1)
tables_list = [
    "Table 1: Group Member Contributions",
    "Table 2: Dataset Feature Description",
    "Table 3: Dataset Statistics (Min, Max, Mean)",
    "Table 4: Model Comparison Metrics",
    "Table 5: Confusion Matrix — Decision Tree",
    "Table 6: Confusion Matrix — Random Forest",
    "Table 7: Confusion Matrix — SVM",
    "Table 8: Input Validation Rules",
    "Table 9: Sample Test Cases",
]
for t in tables_list:
    add_para(t, space_after=2)

doc.add_page_break()

# =====================  6. INTRODUCTION  ===================================
add_heading("5. Introduction", level=1)

add_heading("5.1 Background", level=2)
add_para(
    "The global mobile phone market is one of the largest consumer electronics "
    "sectors, with thousands of devices released annually across vastly different "
    "price segments. Each device is characterized by a unique combination of "
    "hardware specifications — including RAM, battery capacity, camera quality, "
    "processor speed, and screen resolution — that collectively determine its "
    "market positioning and price range."
)

add_heading("5.2 Motivation", level=2)
add_para(
    "Manually categorizing mobile phones into price segments is subjective and "
    "inconsistent. A machine learning-based approach can objectively analyze "
    "hardware features and classify devices into standardized price categories, "
    "benefiting both consumers making purchasing decisions and manufacturers "
    "performing competitive analysis."
)

add_heading("5.3 Importance of Machine Learning", level=2)
add_para(
    "Machine learning algorithms excel at discovering complex, non-linear "
    "relationships between multiple input features and a target variable. For "
    "this multiclass classification problem, ML models can learn the intricate "
    "patterns that map hardware specifications to price ranges — patterns that "
    "would be difficult to capture with simple rule-based systems."
)

add_heading("5.4 Scope", level=2)
add_para(
    "This project covers the complete pipeline from data preprocessing and model "
    "training to web application deployment. The system accepts 11 hardware "
    "features as input, classifies the device into one of four price ranges "
    "(Low, Medium, High, Very High), and presents the result through a modern, "
    "validated web interface built with Flask."
)

doc.add_page_break()

# =====================  7. PROBLEM STATEMENT  ==============================
add_heading("6. Problem Statement", level=1)
add_para(
    "Given a set of mobile phone hardware specifications (battery power, RAM, "
    "internal memory, clock speed, weight, camera megapixels, screen resolution, "
    "and screen dimensions), the objective is to classify the device into one of "
    "four discrete price ranges:"
)
add_bullet("0 — Low Price: Budget segment devices")
add_bullet("1 — Medium Price: Mid-range devices with balanced features")
add_bullet("2 — High Price: Premium devices with strong specifications")
add_bullet("3 — Very High Price: Flagship-tier devices with top-end hardware")

add_para(
    "The expected outcome is a web-based system where users can input device "
    "specifications and receive an instant, accurate price range prediction "
    "powered by a trained machine learning model."
)

doc.add_page_break()

# =====================  8. OBJECTIVES  =====================================
add_heading("7. Objectives of the Project", level=1)

add_heading("7.1 Main Objectives", level=2)
add_bullet("Develop a multiclass classification model to predict mobile phone price ranges.")
add_bullet("Build an interactive web application using Flask for real-time predictions.")

add_heading("7.2 Specific Goals", level=2)
add_bullet("Train and compare at least three ML algorithms (Decision Tree, Random Forest, SVM).")
add_bullet("Select the best-performing model based on accuracy, precision, recall, and F1-score.")
add_bullet("Implement comprehensive two-layer input validation (client-side and server-side).")
add_bullet("Design a modern, responsive UI with glassmorphism styling and color-coded results.")
add_bullet("Provide user-friendly hints showing acceptable input ranges for each feature.")
add_bullet("Document the complete development process in a formal academic report.")

doc.add_page_break()

# =====================  9. DATASET DESCRIPTION  ============================
add_heading("8. Dataset Description", level=1)

add_para("Name: Mobile Price Classification Dataset", bold=True)
add_para("Source: Kaggle (https://www.kaggle.com/datasets/iabhishekofficial/mobile-price-classification)")
add_para("Total Records: 2,000 samples")
add_para("Total Features: 21 (11 selected for the model)")
add_para("Target Variable: price_range (4 classes: 0, 1, 2, 3)")
add_para("")

add_heading("8.1 Input Features Used", level=2)

feature_desc = [
    ("battery_power", "Battery capacity", "mAh"),
    ("ram", "Random Access Memory", "MB"),
    ("int_memory", "Internal storage", "GB"),
    ("clock_speed", "Processor clock speed", "GHz"),
    ("mobile_wt", "Device weight", "grams"),
    ("pc", "Primary (rear) camera resolution", "MP"),
    ("fc", "Front camera resolution", "MP"),
    ("px_height", "Screen pixel resolution height", "pixels"),
    ("px_width", "Screen pixel resolution width", "pixels"),
    ("sc_h", "Screen physical height", "cm"),
    ("sc_w", "Screen physical width", "cm"),
]
make_table(["Feature", "Description", "Unit"], feature_desc)

add_heading("8.2 Dataset Statistics", level=2)
stats = [
    ("battery_power", "501", "1998", "1238.5"),
    ("ram", "256", "3998", "2124.2"),
    ("int_memory", "2", "64", "32.0"),
    ("clock_speed", "0.5", "3.0", "1.5"),
    ("mobile_wt", "80", "200", "140.2"),
    ("pc", "0", "20", "9.9"),
    ("fc", "0", "19", "4.3"),
    ("px_height", "0", "1960", "645.1"),
    ("px_width", "500", "1998", "1251.5"),
    ("sc_h", "5", "19", "12.3"),
    ("sc_w", "0", "18", "5.8"),
]
make_table(["Feature", "Min", "Max", "Mean"], stats)

add_heading("8.3 Target Variable", level=2)
add_para(
    "The target variable price_range has four balanced classes with approximately "
    "500 samples each, ensuring no class imbalance issues."
)
target_dist = [
    ("0 — Low Price", "~500"),
    ("1 — Medium Price", "~500"),
    ("2 — High Price", "~500"),
    ("3 — Very High Price", "~500"),
]
make_table(["Class", "Sample Count"], target_dist)

doc.add_page_break()

# =====================  10. DATA PREPROCESSING  ============================
add_heading("9. Data Preprocessing", level=1)

add_heading("9.1 Handling Missing Values", level=2)
add_para(
    "The dataset was inspected for missing values using pandas. No missing or null "
    "values were found in any of the 21 columns, so no imputation was required."
)

add_heading("9.2 Encoding of Categorical Variables", level=2)
add_para(
    "All 11 selected features are numerical. The dataset does contain binary "
    "categorical features (blue, dual_sim, four_g, three_g, touch_screen, wifi) "
    "but these were not included in the selected feature set, so no encoding was needed."
)

add_heading("9.3 Feature Scaling / Normalization", level=2)
add_para(
    "StandardScaler from scikit-learn was applied to normalize all 11 features to "
    "zero mean and unit variance. This is important because features like RAM "
    "(range: 256–3998) and clock_speed (range: 0.5–3.0) have very different scales. "
    "Without scaling, features with larger magnitudes would dominate the model, "
    "particularly for distance-based algorithms like SVM."
)

add_heading("9.4 Train-Test Split", level=2)
add_para(
    "The dataset was split into 80% training (1,600 samples) and 20% testing "
    "(400 samples) using sklearn's train_test_split with random_state=42 for "
    "reproducibility."
)

add_heading("9.5 Feature Engineering", level=2)
add_para(
    "11 features were selected from the original 21 based on domain knowledge of "
    "mobile hardware specifications. Features like blue (Bluetooth support) and "
    "dual_sim were excluded as they have minimal correlation with price range compared "
    "to specifications like RAM and battery power."
)

doc.add_page_break()

# =====================  11. METHODOLOGY  ===================================
add_heading("10. Methodology", level=1)

add_para("The project follows a structured machine learning pipeline:")

steps = [
    ("Step 1 — Data Loading",
     "The training dataset (train.csv with 2,000 records) is loaded using pandas."),
    ("Step 2 — Feature Selection",
     "11 hardware-related features are selected from the 21 available columns."),
    ("Step 3 — Data Preprocessing",
     "StandardScaler is applied to normalize feature values for consistent model training."),
    ("Step 4 — Train-Test Split",
     "Data is split 80/20 for training and evaluation with a fixed random seed."),
    ("Step 5 — Model Training",
     "Three classifiers (Decision Tree, Random Forest, SVM) are trained on the scaled training data."),
    ("Step 6 — Model Evaluation",
     "Models are compared on accuracy, precision, recall, F1-score, and confusion matrices."),
    ("Step 7 — Model Selection",
     "The best-performing model is selected and saved using joblib (model.pkl, scaler.pkl)."),
    ("Step 8 — Flask Integration",
     "A Flask web server loads the saved model and exposes a /predict API endpoint."),
    ("Step 9 — Frontend Development",
     "An HTML/CSS/JS interface collects user input with validation and displays results."),
    ("Step 10 — Testing & Deployment",
     "The complete pipeline is tested end-to-end with sample inputs and edge cases."),
]
for title, desc in steps:
    add_para(title, bold=True, space_after=2)
    add_para(desc)

doc.add_page_break()

# ================  12. MODELS TRIED  ======================================
add_heading("11. Models Tried for the Project", level=1)

add_heading("11.1 Decision Tree Classifier", level=2)
add_para(
    "Decision Trees split data based on feature thresholds to create a tree-like "
    "structure of decisions. They are highly interpretable and require no feature "
    "scaling. Selected as a baseline model due to simplicity and fast training. "
    "However, they are prone to overfitting without pruning."
)

add_heading("11.2 Random Forest Classifier", level=2)
add_para(
    "Random Forest is an ensemble of multiple decision trees that reduces "
    "overfitting through bagging and random feature selection. It typically "
    "provides higher accuracy than individual decision trees and offers feature "
    "importance rankings. Selected for its robustness and strong performance on "
    "tabular data."
)

add_heading("11.3 Support Vector Machine (SVM)", level=2)
add_para(
    "SVM finds the optimal hyperplane that maximizes the margin between classes. "
    "With the RBF kernel, it can model non-linear decision boundaries effectively. "
    "Selected for its strong theoretical foundation and known effectiveness in "
    "multiclass classification problems. Requires feature scaling for optimal performance."
)

doc.add_page_break()

# ===========  13. MODEL COMPARISON & BEST MODEL  ==========================
add_heading("12. Model Comparison and Best Model Selection", level=1)

add_heading("12.1 Performance Metrics", level=2)
add_para("All three models were evaluated on the same 20% test set (400 samples).")

comparison = [
    ("Decision Tree", "84.25%", "0.8444", "0.8425", "0.8422"),
    ("Random Forest", "90.75%", "0.9093", "0.9075", "0.9082"),
    ("SVM", "91.00%", "0.9122", "0.9100", "0.9104"),
]
make_table(["Model", "Accuracy", "Precision", "Recall", "F1-Score"], comparison)

add_heading("12.2 Confusion Matrices", level=2)

add_para("Decision Tree Confusion Matrix:", bold=True)
make_table(
    ["", "Pred 0", "Pred 1", "Pred 2", "Pred 3"],
    [
        ("Actual 0", "95", "10", "0", "0"),
        ("Actual 1", "5", "79", "7", "0"),
        ("Actual 2", "0", "14", "64", "14"),
        ("Actual 3", "0", "0", "13", "99"),
    ],
)

add_para("Random Forest Confusion Matrix:", bold=True)
make_table(
    ["", "Pred 0", "Pred 1", "Pred 2", "Pred 3"],
    [
        ("Actual 0", "101", "4", "0", "0"),
        ("Actual 1", "2", "82", "7", "0"),
        ("Actual 2", "0", "6", "79", "7"),
        ("Actual 3", "0", "0", "11", "101"),
    ],
)

add_para("SVM Confusion Matrix:", bold=True)
make_table(
    ["", "Pred 0", "Pred 1", "Pred 2", "Pred 3"],
    [
        ("Actual 0", "100", "5", "0", "0"),
        ("Actual 1", "5", "83", "3", "0"),
        ("Actual 2", "0", "12", "77", "3"),
        ("Actual 3", "0", "0", "8", "104"),
    ],
)

add_heading("12.3 Best Model Selection and Justification", level=2)
add_para(
    "While SVM achieves a marginally higher accuracy (91.00%) compared to Random "
    "Forest (90.75%), Random Forest was selected as the final model for the following reasons:"
)
add_bullet("Comparable performance: Only 0.25% difference in accuracy, well within statistical noise.")
add_bullet("Faster inference: Random Forest prediction is significantly faster than SVM, critical for real-time web applications.")
add_bullet("Feature importance: Random Forest provides built-in feature importance rankings useful for interpretability.")
add_bullet("Scalability: Random Forest scales better with increasing data size compared to SVM.")
add_bullet("Stability: Ensemble averaging makes Random Forest more robust to small data perturbations.")
add_bullet("Simplicity: No kernel parameter tuning required, making it easier to maintain in production.")

doc.add_page_break()

# ==============  14. IMPLEMENTATION DETAILS  ===============================
add_heading("13. Implementation Details", level=1)

add_heading("13.1 Programming Language", level=2)
add_para("Python 3.x — chosen for its rich ML ecosystem and web framework support.")

add_heading("13.2 Libraries and Packages", level=2)
libs = [
    ("Flask", "Web framework for API and template rendering"),
    ("scikit-learn", "ML algorithms (Random Forest, Decision Tree, SVM), StandardScaler, train_test_split, metrics"),
    ("pandas", "Data loading and manipulation"),
    ("NumPy", "Numerical operations and array handling"),
    ("joblib", "Model serialization (saving/loading .pkl files)"),
]
make_table(["Package", "Purpose"], libs)

add_heading("13.3 Frontend Technologies", level=2)
fe_libs = [
    ("HTML5", "Page structure and form elements"),
    ("CSS3", "Glassmorphism styling, animations, responsive grid"),
    ("JavaScript (ES6+)", "Client-side validation, async API calls, DOM manipulation"),
    ("Google Fonts (Outfit)", "Typography"),
]
make_table(["Technology", "Purpose"], fe_libs)

add_heading("13.4 Project File Structure", level=2)
add_para(
    "mobile-price-classifier/\n"
    "├── main.py                  # Flask server with prediction API\n"
    "├── train_model.py           # Model training script\n"
    "├── generate_report.py       # Report generation script\n"
    "├── model.pkl                # Trained Random Forest model\n"
    "├── scaler.pkl               # Fitted StandardScaler\n"
    "├── requirements.txt         # Python dependencies\n"
    "├── README.md                # Project documentation\n"
    "├── report.docx              # This report\n"
    "├── dataset/\n"
    "│   └── train.csv            # Training dataset (2000 samples)\n"
    "├── templates/\n"
    "│   └── index.html           # Web interface template\n"
    "└── static/\n"
    "    ├── style.css            # Glassmorphism CSS\n"
    "    └── script.js            # Frontend validation & API logic"
)

doc.add_page_break()

# ==============  15. FRONTEND DESIGN  ======================================
add_heading("14. Frontend Design", level=1)

add_heading("14.1 HTML Page Structure", level=2)
add_para(
    "The application consists of a single-page interface (index.html) rendered by "
    "Flask's Jinja2 template engine. The page contains:"
)
add_bullet("A header section with the project title and subtitle.")
add_bullet("A form with 11 input fields arranged in a responsive CSS Grid (auto-fit, minmax 200px).")
add_bullet("Each input field includes a label, number input with min/max/step attributes, a hint showing the valid range, and an error message span.")
add_bullet("A submit button that triggers the prediction API call.")
add_bullet("A result container (initially hidden) that displays the color-coded prediction and description.")
add_bullet("A footer with technology credits.")

add_heading("14.2 CSS Styling", level=2)
add_para("The UI uses a glassmorphism design system with the following characteristics:")
add_bullet("Dark theme background (#0d1117) with animated gradient blobs for visual depth.")
add_bullet("Glass-effect container with backdrop-filter blur(20px) and subtle borders.")
add_bullet("Cyan accent color (#00f2fe) for interactive elements and focus states.")
add_bullet("Smooth transitions (0.3s ease) on all interactive elements.")
add_bullet("Color-coded result badges: Green (Low), Yellow (Medium), Orange (High), Red (Very High).")
add_bullet("Responsive design with media queries for mobile screens (max-width: 600px).")

add_heading("14.3 Input Form Design", level=2)
add_para(
    "Each input field provides three levels of user guidance: (1) a descriptive label "
    "with the unit, (2) a placeholder example value, and (3) a hint text below "
    "showing the exact valid range. Invalid inputs are highlighted with a red border "
    "and glow effect, with a specific error message displayed below the field."
)

add_heading("14.4 Result Display", level=2)
add_para(
    "The prediction result appears in a pill-shaped badge with color coding that "
    "matches the predicted price tier. Below the badge, a descriptive sentence "
    "explains what the price range means in practical terms (e.g., 'A mid-range "
    "device offering a solid balance of performance and affordability')."
)

doc.add_page_break()

# ==============  16. FLASK INTEGRATION  ====================================
add_heading("15. Flask Integration", level=1)

add_heading("15.1 Route Handling", level=2)
routes = [
    ("GET /", "home()", "Renders index.html template"),
    ("POST /predict", "predict()", "Accepts JSON, validates, predicts, returns result"),
]
make_table(["Route", "Function", "Description"], routes)

add_heading("15.2 Form Input Processing", level=2)
add_para(
    "The /predict endpoint receives input as a JSON payload via a POST request. "
    "Each of the 11 feature values is extracted from the JSON body, validated for "
    "presence, type (must be numeric), and range compliance. If any validation "
    "fails, the server returns HTTP 422 with field-specific error messages."
)

add_heading("15.3 Model Loading and Prediction", level=2)
add_para(
    "At server startup, the trained model (model.pkl) and scaler (scaler.pkl) are "
    "loaded into memory using joblib. For each prediction request, the validated "
    "input features are arranged in the correct order, reshaped into a 1x11 NumPy "
    "array, transformed by the StandardScaler, and passed to the Random Forest model's "
    "predict() method."
)

add_heading("15.4 Output Rendering", level=2)
add_para(
    "The prediction (integer 0–3) is mapped to a human-readable label (e.g., 'High Price') "
    "and a CSS class for color coding. The response is returned as JSON: "
    '{"prediction": "High Price", "css_class": "price-high"}. '
    "The frontend JavaScript updates the DOM to display the result with appropriate styling."
)

doc.add_page_break()

# ==============  17. RESULTS AND DISCUSSION  ===============================
add_heading("16. Results and Discussion", level=1)

add_heading("16.1 Application Output", level=2)
add_para(
    "[Screenshots of the running application to be inserted here showing: "
    "the main form interface, validation errors, and prediction results for different price ranges.]"
)

add_heading("16.2 Sample Inputs and Predicted Outputs", level=2)
samples = [
    ("500, 256, 2, 0.5, 80, 2, 0, 200, 500, 5, 2", "Low Price"),
    ("1200, 2000, 32, 1.8, 150, 10, 5, 800, 1080, 13, 7", "Medium Price"),
    ("1500, 3000, 48, 2.2, 160, 16, 8, 1280, 720, 15, 8", "High Price"),
    ("1800, 3800, 64, 2.8, 180, 20, 15, 1920, 1080, 18, 10", "Very High Price"),
]
make_table(["Input Features (battery, ram, mem, clock, wt, pc, fc, pxH, pxW, scH, scW)", "Predicted Range"], samples)

add_heading("16.3 Discussion", level=2)
add_para(
    "The Random Forest model achieves 90.75% accuracy on the test set, correctly "
    "classifying the majority of devices across all four price categories. The model "
    "performs best on extreme categories (Low and Very High) where hardware specifications "
    "are most distinct, with slightly lower performance on Medium and High categories "
    "where feature overlap is greater."
)
add_para(
    "RAM and battery power emerge as the most influential features for price classification, "
    "which aligns with real-world consumer expectations. The two-layer validation system "
    "ensures that only valid inputs reach the model, preventing nonsensical predictions."
)

doc.add_page_break()

# ==============  18. TESTING  ==============================================
add_heading("17. Testing", level=1)

add_heading("17.1 Functional Testing", level=2)
add_para(
    "The application was tested for correct functionality across the following areas:"
)
add_bullet("Form submission with valid inputs produces correct predictions.")
add_bullet("Empty form submission triggers client-side validation errors on all fields.")
add_bullet("Out-of-range values display appropriate error messages with valid ranges.")
add_bullet("Server-side validation catches tampered requests (e.g., via curl or Postman).")
add_bullet("Result display shows correct color coding for each price category.")
add_bullet("Page is responsive and functional on mobile screen sizes.")

add_heading("17.2 Sample Test Cases", level=2)
test_cases = [
    ("All fields empty", "Submit", "Validation errors on all 11 fields", "Pass"),
    ("battery_power = 99999", "Submit", "'Must be between 500 and 2000'", "Pass"),
    ("ram = -100", "Submit", "'Must be between 256 and 4000'", "Pass"),
    ("clock_speed = abc", "Submit", "'Please enter a valid number'", "Pass"),
    ("All valid (low specs)", "Submit", "Prediction: Low Price (green)", "Pass"),
    ("All valid (high specs)", "Submit", "Prediction: Very High Price (red)", "Pass"),
    ("Server receives empty JSON", "API call", "HTTP 422 with error list", "Pass"),
    ("Correct input then fix error", "Type", "Error clears on input change", "Pass"),
]
make_table(["Test Case", "Action", "Expected Result", "Status"], test_cases)

add_heading("17.3 Prediction Flow Verification", level=2)
add_para(
    "The end-to-end prediction flow was verified by comparing web application outputs "
    "with direct Python model predictions using the same inputs. All results matched, "
    "confirming that the Flask integration, scaling, and prediction pipeline are functioning correctly."
)

doc.add_page_break()

# ==============  19. CHALLENGES FACED  =====================================
add_heading("18. Challenges Faced", level=1)

add_heading("18.1 Data-Related Issues", level=2)
add_bullet("Feature selection: Choosing the most relevant 11 features from 21 required domain analysis and experimentation to find the subset that balances accuracy with practical input collection.")
add_bullet("Feature scale disparity: RAM values (256–3998) vs clock_speed (0.5–3.0) required StandardScaler to prevent magnitude-biased predictions.")

add_heading("18.2 Model Performance Issues", level=2)
add_bullet("Decision Tree overfitting: Initial Decision Tree models memorized training data; Random Forest's ensemble approach mitigated this.")
add_bullet("SVM training time: SVM with RBF kernel was significantly slower to train than tree-based models, making iterative experimentation time-consuming.")
add_bullet("Mid-range confusion: All models showed some difficulty distinguishing between Medium and High price categories due to overlapping feature distributions.")

add_heading("18.3 Flask Integration Issues", level=2)
add_bullet("Feature ordering: The feature array order in the Flask API must exactly match the order used during model training; a mismatch caused incorrect predictions during initial integration.")
add_bullet("Scaler consistency: The same fitted scaler must be used at prediction time; re-fitting on new data would shift predictions.")

add_heading("18.4 Frontend-Backend Communication", level=2)
add_bullet("JSON data types: Input values from HTML forms are strings by default; they must be converted to floats before passing to the model.")
add_bullet("Error handling: Designing a unified error response format that both the server-side validation and client-side JavaScript could handle consistently required careful API design.")
add_bullet("CORS and content-type headers had to be configured correctly for the fetch API to communicate with Flask.")

doc.add_page_break()

# ==============  20. CONCLUSION  ===========================================
add_heading("19. Conclusion", level=1)

add_heading("19.1 Summary of Work", level=2)
add_para(
    "This project successfully developed a complete web-based Mobile Price Range "
    "Classifier that predicts whether a mobile phone falls into the Low, Medium, "
    "High, or Very High price category based on 11 hardware specifications. The "
    "system integrates a machine learning model with a Flask web application, providing "
    "real-time predictions through a modern, user-friendly interface."
)

add_heading("19.2 Achievement of Objectives", level=2)
add_bullet("Trained and compared three ML models: Decision Tree (84.25%), Random Forest (90.75%), and SVM (91.00%).")
add_bullet("Selected Random Forest as the production model for its balance of accuracy, speed, and interpretability.")
add_bullet("Built a complete Flask web application with validated input forms and color-coded results.")
add_bullet("Implemented two-layer validation (client + server) with user-friendly range hints and error messages.")
add_bullet("Designed a responsive glassmorphism UI that works across desktop and mobile devices.")
add_bullet("Documented the entire development process in this comprehensive report.")

add_heading("19.3 Learning Outcomes", level=2)
add_bullet("Gained hands-on experience with the complete ML pipeline: data preprocessing, model training, evaluation, and deployment.")
add_bullet("Learned to build and integrate ML models with web frameworks (Flask).")
add_bullet("Understood the importance of input validation in production ML systems.")
add_bullet("Practiced comparative analysis of ML algorithms using standard evaluation metrics.")
add_bullet("Developed skills in frontend design, responsive CSS, and client-server communication.")

# =====================  SAVE  ==============================================
doc.save("report.docx")
print("report.docx generated successfully!")
